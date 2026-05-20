"""Shared helpers for client guide DOCX generation."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

BOT_USERNAME = "thali_manager_bot"
BOT_URL = f"https://t.me/{BOT_USERNAME}"


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    return doc


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(
        "Инструкция по Telegram-боту\nресторана «Тхали и Карри»"
    )
    tr.bold = True
    tr.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Для руководителя и сотрудников\nОбновлено: май 2026")
    sr.font.size = Pt(11)

    doc.add_paragraph()

    bot_p = doc.add_paragraph()
    bot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br1 = bot_p.add_run("Бот в Telegram: ")
    br1.font.size = Pt(12)
    br1.bold = True
    br2 = bot_p.add_run(f"@{BOT_USERNAME}")
    br2.font.size = Pt(12)
    br2.bold = True
    br3 = bot_p.add_run(f"\n{BOT_URL}")
    br3.font.size = Pt(11)

    hint = doc.add_paragraph()
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hint.add_run(
        "Откройте ссылку на телефоне или найдите бота по имени в поиске Telegram."
    )
    hr.font.size = Pt(10)
    hr.italic = True

    doc.add_paragraph()


def add_footer(doc: Document) -> None:
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        f"Бот: @{BOT_USERNAME}  |  {BOT_URL}\n"
        "Техническая поддержка и настройка таблицы - у вашего разработчика бота."
    )
    fr.font.size = Pt(9)
    fr.italic = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
